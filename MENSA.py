import sys
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QLabel, QLineEdit, QPushButton, 
    QMessageBox, QProgressBar, QListWidget, QInputDialog, QTableWidget, QTableWidgetItem,
    QHeaderView, QFormLayout, QStackedWidget, QWidget, QVBoxLayout, QLabel, QTableWidget, QTableWidgetItem, QHBoxLayout, 
                             QInputDialog, QComboBox, QPushButton, QFileDialog, QMessageBox, QFormLayout, QLineEdit
)
from PyQt6.QtCore import Qt, QTimer, QSize
from PyQt6.QtGui import QFont, QPixmap, QIcon, QColor, QBrush
import mysql.connector
import matplotlib.pyplot as plt
import seaborn as sns
from fpdf import FPDF
from docx import Document

from PyQt6.QtPrintSupport import QPrinter, QPrintDialog  # For PDF export
import pandas as pd  # For MS Excel export


class LoadingScreen(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle('MENSA')
        self.setFixedSize(600, 400)
        self.setWindowIcon(QIcon("title_logo.png"))
        self.setStyleSheet("background-color: #176276;")
        layout = QVBoxLayout()
        self.setLayout(layout)
        
        logo = QLabel(self)
        pixmap = QPixmap("load_logo.png")
        logo.setPixmap(pixmap.scaled(300, 300, Qt.AspectRatioMode.KeepAspectRatio))
        logo.setAlignment(Qt.AlignmentFlag.AlignCenter)

               
        self.loading_bar = QProgressBar()
        self.loading_bar.setStyleSheet("QProgressBar {background-color: white; border-radius: 10px; text-align: center; color: white}"
                                       "QProgressBar::chunk {background-color: #176276;}")
        self.loading_bar.setRange(0, 100)
        self.loading_bar.setValue(0)

        layout.addWidget(logo)
        layout.addWidget(self.loading_bar)
        layout.setAlignment(Qt.AlignmentFlag.AlignCenter)
        
        self.timer = QTimer()
        self.timer.timeout.connect(self.update_progress)
        self.timer.start(50)

    def update_progress(self):
        value = self.loading_bar.value() + 1
        self.loading_bar.setValue(value)
        if value >= 100:
            self.timer.stop()
            self.close()
            self.open_login_window()

    def open_login_window(self):
        self.login_window = LoginWindow()
        self.login_window.show()

class LoginWindow(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle('MENSA')
        self.setFixedSize(600, 500)
        self.setWindowIcon(QIcon("title_logo.png"))
        self.setStyleSheet("background-color: white;")
        layout = QVBoxLayout()
        self.setLayout(layout)
        
        logo = QLabel(self)
        pixmap = QPixmap("log_logo.png")
        logo.setPixmap(pixmap.scaled(350, 350, Qt.AspectRatioMode.KeepAspectRatio))
        logo.setAlignment(Qt.AlignmentFlag.AlignCenter)
        
             
        self.username_input = QLineEdit(self)
        self.username_input.setPlaceholderText("USERNAME")
        self.username_input.setFont(QFont('Century Gothic', 14))
        self.username_input.setAlignment(Qt.AlignmentFlag.AlignCenter)
        
        self.password_input = QLineEdit(self)
        self.password_input.setPlaceholderText("PASSWORD")
        self.password_input.setFont(QFont('Century Gothic', 14))
        self.password_input.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.password_input.setEchoMode(QLineEdit.EchoMode.Password)
        
        login_button = QPushButton("Login", self)
        login_button.setFont(QFont('Century Gothic', 16))
        login_button.setStyleSheet("background-color: #176276; color: white;")
        login_button.clicked.connect(self.handle_login)

        layout.addWidget(logo)
        layout.addWidget(self.username_input)
        layout.addWidget(self.password_input)
        layout.addWidget(login_button)
        layout.setAlignment(Qt.AlignmentFlag.AlignCenter)

    def handle_login(self):
        username = self.username_input.text()
        password = self.password_input.text()
        self.connecting_window = ConnectingWindow(username, password)
        self.connecting_window.show()
        self.close()

class ConnectingWindow(QWidget):
    def __init__(self, username, password):
        super().__init__()
        self.setWindowTitle('MENSA')
        self.setFixedSize(600, 400)
        self.setWindowIcon(QIcon("title_logo.png"))
        self.setStyleSheet("background-color: #176276;")
        layout = QVBoxLayout()
        self.setLayout(layout)
        
        self.status_label = QLabel("Connecting to MySQL Database....")
        self.status_label.setFont(QFont('Century Gothic', 25))
        self.status_label.setStyleSheet("color: white;")
        self.status_label.setAlignment(Qt.AlignmentFlag.AlignCenter)

        layout.addWidget(self.status_label)
        layout.setAlignment(Qt.AlignmentFlag.AlignCenter)

        self.timer = QTimer()
        self.timer.timeout.connect(lambda: self.check_connection(username, password))
        self.timer.start(2000)

    def check_connection(self, username, password):
        self.timer.stop()
        try:
            self.connection = mysql.connector.connect(
                host="localhost",
                user=username,
                password=password
            )
            self.status_label.setText("Connection Successful")
            QTimer.singleShot(2000, self.open_main_app)
        except mysql.connector.Error as err:
            self.status_label.setText("Unable to connect with the \nMySQL Database Server.\nPlease re-check your \nlogin credentials")
            QTimer.singleShot(2000, self.open_login_window)

    def open_login_window(self):
        self.login_window = LoginWindow()
        self.login_window.show()
        self.close()

    def open_main_app(self):
        self.main_app = MainApp(self.connection)
        self.main_app.show()
        self.close()

class MainApp(QMainWindow):
    def __init__(self, connection):
        super().__init__()
        self.connection = connection
        self.setWindowTitle('MENSA - MySQL Enhanced for Navigational Speed and Administration')
        self.setWindowIcon(QIcon("title_logo.png"))
        self.setStyleSheet("background-color: white; color: black;")
        self.showMaximized()
        self.initUI()

        # Flags for Dark Mode and Colorblind Mode
        self.dark_mode_enabled = False
        self.colorblind_mode_enabled = False

    def initUI(self):
        """
        Initialize the central widget and setup the UI.
        """
        self.central_widget = QStackedWidget()
        self.setCentralWidget(self.central_widget)

        # Initialize and add DatabaseArena as the first screen
        self.db_arena = DatabaseArena(self.connection, self)
        self.central_widget.addWidget(self.db_arena)

        # Create the top menu
        self.create_menu()

    def create_menu(self):
        self.menu = self.menuBar()

                
        self.menu.setStyleSheet("background-color: #176276; color: #F9F4E0; font: 12pt 'Corbel';")

        # Logout
        logout_action = self.menu.addAction("Logout")
        logout_action.triggered.connect(self.logout)

        # Refresh
        refresh_action = self.menu.addAction("Refresh")
        refresh_action.triggered.connect(self.refresh)

        # Help
        help_action = self.menu.addAction("Help")
        help_action.triggered.connect(self.show_help)

        # Reach Us
        reach_us_menu = self.menu.addMenu("Reach Us")
        reach_us_action = reach_us_menu.addAction("Reach Us")
        reach_us_action.triggered.connect(self.open_reach_us)
        feedback_action = reach_us_menu.addAction("Feedback")
        feedback_action.triggered.connect(self.open_feedback)

        # Dark Mode Toggle
        dark_mode_action = self.menu.addAction("Dark Mode Toggle")
        dark_mode_action.triggered.connect(self.toggle_dark_mode)

        # Colorblind Mode
        colorblind_mode_action = self.menu.addAction("Colorblind Mode")
        colorblind_mode_action.triggered.connect(self.activate_colorblind_mode)

        no_action = self.menu.addAction("                                                                                                                                                                                                                                    ")
        no_action.triggered.connect(self.no)

     
        # MENSA Tools - Right Aligned
        mensa_tools_menu = self.menu.addMenu("MENSA TOOLS")

    
        mensa_tools_menu.setStyleSheet("background-color: #176276; color: #F9F4E0")
        mensa_effingo_action = mensa_tools_menu.addAction("MENSA EFFINGO")
        mensa_effingo_action.triggered.connect(self.mensa_effingo)

        # Spacer to move MENSA TOOLS to the rightmost corner
        spacer = self.menu.addAction("")
        spacer.setEnabled(False)
        spacer.setVisible(False)
        self.menu.insertSeparator(spacer)

    def mensa_effingo(self):
        """
        MENSA EFFINGO Functionality: Duplicate a table from one database to another.
        """
        try:
            # Step 1: Select Source Database
            cursor = self.connection.cursor()
            cursor.execute("SHOW DATABASES")
            databases = [db[0] for db in cursor.fetchall()]
            cursor.close()

            if not databases:
                QMessageBox.warning(self, "No Databases Found", "No databases available for copying tables.")
                return

            source_db, ok = QInputDialog.getItem(self, "MENSA EFFINGO", "Select the source database:", databases, 0, False)
            if not ok or not source_db:
                return

            # Step 2: Select Table to Copy
            cursor = self.connection.cursor()
            cursor.execute(f"USE {source_db}")
            cursor.execute("SHOW TABLES")
            tables = [table[0] for table in cursor.fetchall()]
            cursor.close()

            if not tables:
                QMessageBox.warning(self, "No Tables Found", f"No tables found in the database '{source_db}'.")
                return

            table_to_copy, ok = QInputDialog.getItem(
                self, "MENSA EFFINGO", f"Select the table to copy from '{source_db}':", tables, 0, False
            )
            if not ok or not table_to_copy:
                return

            # Step 3: Select Target Database
            target_db, ok = QInputDialog.getItem(self, "MENSA EFFINGO", "Select the target database:", databases, 0, False)
            if not ok or not target_db:
                return

            # Step 4: Name the Duplicate Table
            new_table_name, ok = QInputDialog.getText(
                self, "MENSA EFFINGO", f"Enter the name for the duplicate table in '{target_db}':"
            )
            if not ok or not new_table_name:
                return

            # Step 5: Check if the table name already exists in the target database
            cursor = self.connection.cursor()
            cursor.execute(f"USE {target_db}")
            cursor.execute("SHOW TABLES")
            existing_tables = [table[0] for table in cursor.fetchall()]

            if new_table_name in existing_tables:
                QMessageBox.warning(
                    self,
                    "Duplicate Table Name",
                    f"The table name '{new_table_name}' already exists in the database '{target_db}'. Please try again.",
                )
                return

            # Step 6: Create Duplicate Table
            cursor.execute(f"USE {source_db}")
            cursor.execute(f"CREATE TABLE {target_db}.{new_table_name} LIKE {source_db}.{table_to_copy}")
            cursor.execute(f"INSERT INTO {target_db}.{new_table_name} SELECT * FROM {source_db}.{table_to_copy}")
            self.connection.commit()
            cursor.close()

            QMessageBox.information(
                self, "Success", f"Table '{table_to_copy}' from '{source_db}' was successfully copied to '{target_db}' as '{new_table_name}'."
            )

        except Exception as e:
            QMessageBox.critical(self, "Error", f"An error occurred during the operation: {str(e)}")

    def logout(self):
        """
        Handle logout functionality.
        """
        reply = QMessageBox.question(self, 'Logout', 'Are you sure you want to logout?',
                                     QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                                     QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            self.close()
            self.login_window = LoginWindow()
            self.login_window.show()

    def refresh(self):
        """
        Refresh the current arena.
        """
        self.db_arena.update_databases()

    def no(self):
        
        pass
    
    def open_reach_us(self):
        """
        Open the 'Reach Us' URL in a web browser.
        """
        import webbrowser
        webbrowser.open("https://heyguys.my.canva.site/mensa")

    def open_feedback(self):
        """
        Open the feedback form URL in a web browser.
        """
        import webbrowser
        webbrowser.open("https://forms.gle/fnHgC6gdZDQfvFeH6")

    def show_help(self):
        """
        Display a help window with details about the app.
        """
        help_window = QMessageBox(self)
        help_window.setWindowTitle("About MENSA")
        help_window.setText(
            "MENSA: MySQL Enhanced for Navigational Speed and Administration.\n"
            "A versatile and intuitive GUI tool for database management.\n"
            "Explore databases, manage tables, and visualize data seamlessly!"
        )
        help_window.setStandardButtons(QMessageBox.StandardButton.Ok)
        help_window.exec()

    def toggle_dark_mode(self):
        """
        Toggles between dark mode and normal mode.
        """
        if not self.dark_mode_enabled:
            self.setStyleSheet("background-color: black; color: white;")
            self.menu.setStyleSheet("background-color: black; color: white; font: 12pt 'Corbel';")
            self.dark_mode_enabled = True
        else:
            self.setStyleSheet("background-color: white; color: black;")
            self.menu.setStyleSheet("background-color: #176276; color: #F9F4E0; font: 12pt 'Corbel';")
            self.dark_mode_enabled = False

    def activate_colorblind_mode(self):
        """
        Activates colorblind mode with user-selected options.
        """
        # Show dialog for colorblindness type selection
        options = ["Protanopia (Red-Blind)", "Deuteranopia (Green-Blind)", "Tritanopia (Blue-Blind)", "Monochromacy"]
        colorblind_type, ok = QInputDialog.getItem(self, "Colorblind Mode",
                                                   "Select the type of colorblindness:", options, 0, False)
        if ok and colorblind_type:
            # Apply relevant changes to the color palette
            if colorblind_type == "Protanopia (Red-Blind)":
                self.setStyleSheet("background-color: #a9a9a9; color: #4b0082;")
            elif colorblind_type == "Deuteranopia (Green-Blind)":
                self.setStyleSheet("background-color: #d3d3d3; color: #8b0000;")
            elif colorblind_type == "Tritanopia (Blue-Blind)":
                self.setStyleSheet("background-color: #f5deb3; color: #483d8b;")
            elif colorblind_type == "Monochromacy":
                self.setStyleSheet("background-color: #f0f0f0; color: #000000;")

            # Inform user about the activation
            QMessageBox.information(self, "Colorblind Mode",
                                    "SEE THE WORLD YOUR WAY!!!!\nHope you Find MENSA useful :)")

    def switch_widget(self, widget):
        """
        Switch to the given widget in the QStackedWidget.
        :param widget: The widget to switch to.
        """
        self.central_widget.setCurrentWidget(widget)




class DatabaseArena(QWidget):
    def __init__(self, connection, main_window):
        super().__init__()
        self.connection = connection
        self.main_window = main_window
        self.setStyleSheet("background-color: white; color: black;")  # Default theme
        layout = QVBoxLayout()
        self.setLayout(layout)

        # Header        
        header = QLabel("DATABASES")
        header.setFont(QFont('High Tower Text', 26, QFont.Weight.Bold)) 
        header.setStyleSheet("color: #176276;")
        header.setAlignment(Qt.AlignmentFlag.AlignCenter)

        # Toolbar and Database List
        self.tool_bar = DatabaseToolBar(self.connection, self)
        self.database_list = QListWidget()
        self.database_list.setFont(QFont('Gadugi', 14))
        self.database_list.itemDoubleClicked.connect(self.open_tables_arena)

        # Add to Layout
        layout.addWidget(header)
        layout.addWidget(self.tool_bar)
        layout.addWidget(self.database_list)
        layout.setAlignment(Qt.AlignmentFlag.AlignTop)

        self.update_databases()

    def update_databases(self):
        """
        Fetch and display all databases from the MySQL server.
        """
        self.database_list.clear()
        try:
            databases = self.fetch_databases()
            for db in databases:
                self.database_list.addItem(db[0])
        except Exception as e:
            QMessageBox.critical(self, "Database Error", f"Failed to fetch databases: {e}")

    def fetch_databases(self):
        """
        Execute the query to fetch all database names.
        """
        cursor = self.connection.cursor()
        cursor.execute("SHOW DATABASES")
        databases = cursor.fetchall()
        cursor.close()
        return databases

    def open_tables_arena(self, item):
        """
        Transition to the TablesArena for the selected database.
        """
        try:
            database_name = item.text()
            cursor = self.connection.cursor()
            cursor.execute(f"USE {database_name}")  # Switch to the selected database
            cursor.close()

            # Open the TablesArena
            self.tables_arena = TablesArena(self.connection, database_name, self.main_window)
            self.main_window.central_widget.addWidget(self.tables_arena)
            self.main_window.switch_widget(self.tables_arena)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to open tables for {item.text()}: {e}")

    def apply_dark_mode(self, enable):
        """
        Changes the theme to dark mode or reverts to default.
        """
        if enable:
            self.setStyleSheet("background-color: black; color: white;")
            self.database_list.setStyleSheet("color: white; background-color: black;")
        else:
            self.setStyleSheet("background-color: white; color: black;")
            self.database_list.setStyleSheet("")


class DatabaseToolBar(QWidget):
    def __init__(self, connection, parent):
        super().__init__()
        self.connection = connection
        self.parent = parent
        self.setStyleSheet("background-color: white; border: 1px solid navy;")
        layout = QHBoxLayout()
        self.setLayout(layout)

        # Export Button
        export_db_btn = QPushButton("EXPORT DATABASE LIST")
        export_db_btn.setFont(QFont('Candara', 14))
        export_db_btn.setStyleSheet("background-color: #176276; color: #F9F4E0;")
        export_db_btn.clicked.connect(self.export_databases)

        # Sort Button
        sort_btn = QPushButton("SORT DATABASES")
        sort_btn.setFont(QFont('Candara', 14))
        sort_btn.setStyleSheet("background-color: #176276; color: #F9F4E0;")
        sort_btn.clicked.connect(self.sort_databases)

        # Search Button
        search_btn = QPushButton("SEARCH DATABASE")
        search_btn.setFont(QFont('Candara', 14))
        search_btn.setStyleSheet("background-color: #176276; color: #F9F4E0;")
        search_btn.clicked.connect(self.search_database)

        # Reset Button
        self.reset_btn = QPushButton("RESET")
        self.reset_btn.setFont(QFont('Candara', 14))
        self.reset_btn.setStyleSheet("background-color: #176276; color: #F9F4E0;")
        self.reset_btn.clicked.connect(self.reset_databases)
        self.reset_btn.setVisible(False)  # Hidden by default

        # Go Back Button
        go_back_btn = QPushButton("GO BACK")
        go_back_btn.setFont(QFont('Candara', 14))
        go_back_btn.setStyleSheet("background-color: #176276; color: #F9F4E0;")
        go_back_btn.clicked.connect(self.go_back)

        # Add Buttons to Layout
        layout.addWidget(export_db_btn)
        layout.addWidget(sort_btn)
        layout.addWidget(search_btn)
        layout.addWidget(self.reset_btn)
        layout.addWidget(go_back_btn)
        layout.setAlignment(Qt.AlignmentFlag.AlignTop)

    def export_databases(self):
        """
        Export the list of all databases to a file.
        """
        try:
            # Fetch all databases
            cursor = self.connection.cursor()
            cursor.execute("SHOW DATABASES")
            databases = [db[0] for db in cursor.fetchall()]
            cursor.close()

            if not databases:
                QMessageBox.warning(self, "No Databases Found", "No databases found to export.")
                return

            # Choose export format
            options = ["Excel", "CSV", "PDF", "Word"]
            file_type, ok = QInputDialog.getItem(self, "Export Database List", "Select file format:", options, 0, False)
            if ok and file_type:
                file_path, _ = QFileDialog.getSaveFileName(
                    self, "Save File", f"Database_List.{file_type.lower()}", f"{file_type} Files (*.{file_type.lower()})"
                )
                if file_path:
                    # Perform export based on file type
                    self.perform_export(databases, file_path, file_type)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to export database list: {str(e)}")

    def perform_export(self, databases, file_path, file_type):
        """
        Perform the actual export of databases to the selected file type.
        """
        try:
            if file_type == "Excel":
                import pandas as pd
                pd.DataFrame(databases, columns=["Databases"]).to_excel(file_path, index=False)
            elif file_type == "CSV":
                import pandas as pd
                pd.DataFrame(databases, columns=["Databases"]).to_csv(file_path, index=False)
            elif file_type == "PDF":
                from fpdf import FPDF
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("Arial", size=12)
                pdf.cell(200, 10, txt="Database List", ln=1, align="C")
                for db in databases:
                    pdf.cell(200, 10, txt=db, ln=1)
                pdf.output(file_path)
            elif file_type == "Word":
                from docx import Document
                doc = Document()
                doc.add_heading("Database List", level=1)
                for db in databases:
                    doc.add_paragraph(db)
                doc.save(file_path)
            QMessageBox.information(self, "Export Successful", f"Database list exported successfully to {file_path}!")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to export: {str(e)}")

    def sort_databases(self):
        """
        Sort the list of databases based on selected criteria.
        """
        sort_criteria, ok = QInputDialog.getItem(
            self, 'Sort Databases', 'How do you want to sort the databases?',
            ['Ascending', 'Descending', 'By Name Length (Small to Big)', 'By Name Length (Big to Small)'], 0, False
        )
        if ok:
            databases = [self.parent.database_list.item(i).text() for i in range(self.parent.database_list.count())]
            if sort_criteria == 'Ascending':
                databases.sort()
            elif sort_criteria == 'Descending':
                databases.sort(reverse=True)
            elif sort_criteria == 'By Name Length (Small to Big)':
                databases.sort(key=len)
            elif sort_criteria == 'By Name Length (Big to Small)':
                databases.sort(key=len, reverse=True)

            self.parent.database_list.clear()
            for db in databases:
                self.parent.database_list.addItem(db)

            # Show Reset Button
            self.reset_btn.setVisible(True)

    def search_database(self):
        """
        Search for a database by keyword.
        """
        keyword, ok = QInputDialog.getText(self, "Search Database", "Enter the keyword to search:")
        if ok and keyword:
            matching_databases = [
                self.parent.database_list.item(i).text()
                for i in range(self.parent.database_list.count())
                if keyword.lower() in self.parent.database_list.item(i).text().lower()
            ]
            self.parent.database_list.clear()
            if matching_databases:
                for db in matching_databases:
                    self.parent.database_list.addItem(db)
            else:
                QMessageBox.information(self, "No Results", f"No databases found matching '{keyword}'.")

            # Show Reset Button
            self.reset_btn.setVisible(True)

    def reset_databases(self):
        """
        Reset the database list to its original state.
        """
        self.parent.update_databases()
        self.reset_btn.setVisible(False)

    def go_back(self):
        """
        Go back to the main Database Arena.
        """
        self.parent.main_window.switch_widget(self.parent.main_window.db_arena)

class TablesArena(QWidget):
    def __init__(self, connection, database_name, main_window):
        super().__init__()
        self.connection = connection
        self.database_name = database_name
        self.main_window = main_window

        # Attributes to manage the table list
        self.original_table_list = []  # Store the original list of tables
        self.filtered = False  # Track whether the list is modified (sorted/searched)

        self.setWindowTitle(f'Tables in Database - {database_name}')
        self.setStyleSheet("background-color: white; color: black;")
        layout = QVBoxLayout()
        self.setLayout(layout)

        # Header
        header = QLabel(f"TABLES IN DATABASE - {database_name}")
        header.setFont(QFont('High Tower Text', 26, QFont.Weight.Bold))
        header.setStyleSheet("color: #176276;")
        header.setAlignment(Qt.AlignmentFlag.AlignCenter)

        # Toolbar and Table List
        self.tool_bar = TablesToolBar(self.connection, self)
        self.table_list = QListWidget()
        self.table_list.setFont(QFont('Gadugi', 14))
        self.table_list.itemDoubleClicked.connect(self.open_table_viewer)

        # Add Widgets to Layout
        layout.addWidget(header)
        layout.addWidget(self.tool_bar)
        layout.addWidget(self.table_list)
        layout.setAlignment(Qt.AlignmentFlag.AlignTop)

        # Populate the table list
        self.update_tables()

    def update_tables(self):
        """
        Fetch and display the original table list from the database.
        """
        try:
            cursor = self.connection.cursor()
            cursor.execute(f"USE {self.database_name}")
            cursor.execute("SHOW TABLES")
            self.original_table_list = [table[0] for table in cursor.fetchall()]
            cursor.close()

            # Update the displayed list
            self.refresh_table_list(self.original_table_list)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to fetch tables: {str(e)}")

    def refresh_table_list(self, table_list):
        """
        Refresh the displayed table list with the provided data.
        """
        self.table_list.clear()
        for table in table_list:
            self.table_list.addItem(table)

    def restore_original_tables(self):
        """
        Restore the original table list when the refresh button is clicked.
        """
        if self.filtered:  # Only reset if the list was modified
            self.refresh_table_list(self.original_table_list)
            self.filtered = False

    def open_table_viewer(self, item):
        """
        Open the content of the selected table in the TableViewer.
        """
        try:
            self.table_viewer = TableViewer(self.connection, self.database_name, item.text(), self.main_window)
            self.main_window.central_widget.addWidget(self.table_viewer)
            self.main_window.switch_widget(self.table_viewer)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to open table content: {str(e)}")

class TablesToolBar(QWidget):
    def __init__(self, connection, parent):
        super().__init__()
        self.connection = connection
        self.parent = parent
        self.setStyleSheet("background-color: white; border: 1px solid navy;")
        layout = QHBoxLayout()
        self.setLayout(layout)

        # Sort Button
        sort_table_btn = QPushButton("SORT TABLES")
        sort_table_btn.setFont(QFont('Candara', 14))
        sort_table_btn.setStyleSheet("background-color: #176276; color: #F9F4E0;")
        sort_table_btn.clicked.connect(self.sort_tables)

        # View Table Structure Button
        view_structure_btn = QPushButton("VIEW TABLE STRUCTURE")
        view_structure_btn.setFont(QFont('Candara', 14))
        view_structure_btn.setStyleSheet("background-color: #176276; color: #F9F4E0;")
        view_structure_btn.clicked.connect(self.view_table_structure)

        # Search Button
        search_table_btn = QPushButton("SEARCH TABLE")
        search_table_btn.setFont(QFont('Candara', 14))
        search_table_btn.setStyleSheet("background-color: #176276; color: #F9F4E0;")
        search_table_btn.clicked.connect(self.search_tables)

        # Go Back Button
        go_back_btn = QPushButton("GO BACK")
        go_back_btn.setFont(QFont('Candara', 14))
        go_back_btn.setStyleSheet("background-color: #176276; color: #F9F4E0;")
        go_back_btn.clicked.connect(self.go_back)

        # Add Buttons to Layout
        layout.addWidget(sort_table_btn)
        layout.addWidget(view_structure_btn)
        layout.addWidget(search_table_btn)
        layout.addWidget(go_back_btn)
        layout.setAlignment(Qt.AlignmentFlag.AlignTop)

    def sort_tables(self):
        """
        Sort the tables based on user-selected criteria.
        """
        sort_criteria, ok = QInputDialog.getItem(
            self, 'Sort Tables', 'How do you want to sort the tables?',
            ['Ascending', 'Descending', 'By Name Length (Small to Big)', 'By Name Length (Big to Small)'], 0, False
        )
        if ok:
            tables = self.parent.original_table_list.copy()
            if sort_criteria == 'Ascending':
                tables.sort()
            elif sort_criteria == 'Descending':
                tables.sort(reverse=True)
            elif sort_criteria == 'By Name Length (Small to Big)':
                tables.sort(key=len)
            elif sort_criteria == 'By Name Length (Big to Small)':
                tables.sort(key=len, reverse=True)

            self.parent.refresh_table_list(tables)
            self.parent.filtered = True  # Mark the list as modified

    def view_table_structure(self):
        """
        View the structure of a selected table.
        """
        selected_table, ok = QInputDialog.getItem(
            self, "View Table Structure", "Select a table to view its structure:", self.parent.original_table_list, 0, False
        )
        if ok and selected_table:
            confirm = QMessageBox.question(
                self, "Confirm Action", f"Do you want to view the structure of the table '{selected_table}'?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No, QMessageBox.StandardButton.No
            )
            if confirm == QMessageBox.StandardButton.Yes:
                try:
                    cursor = self.connection.cursor()
                    cursor.execute(f"USE {self.parent.database_name}")
                    cursor.execute(f"DESC {selected_table}")
                    structure_data = cursor.fetchall()
                    columns = [desc[0] for desc in cursor.description]
                    cursor.close()

                    # Open Structure Viewer Window
                    self.structure_window = StructureViewer(self.connection, selected_table, structure_data, columns)
                    self.structure_window.show()
                except Exception as e:
                    QMessageBox.critical(self, "Error", f"Failed to fetch table structure: {str(e)}")

    def search_tables(self):
        """
        Search for tables matching a keyword.
        """
        keyword, ok = QInputDialog.getText(self, "Search Tables", "Enter the keyword to search:")
        if ok and keyword:
            matching_tables = [
                table for table in self.parent.original_table_list
                if keyword.lower() in table.lower()
            ]
            self.parent.refresh_table_list(matching_tables)
            self.parent.filtered = True  # Mark the list as modified

    def go_back(self):
        """
        Navigate back to the DatabaseArena.
        """
        self.parent.main_window.switch_widget(self.parent.main_window.db_arena)

class StructureViewer(QWidget):
    def __init__(self, connection, table_name, structure_data, columns):
        super().__init__()
        self.connection = connection
        self.table_name = table_name
        self.structure_data = structure_data
        self.columns = columns

        self.setWindowTitle(f"Structure of Table '{table_name}'")
        self.setStyleSheet("background-color: white; color: black;")
        layout = QVBoxLayout()
        self.setLayout(layout)

        

        # Header
        header = QLabel(f"STRUCTURE OF TABLE: {table_name}")
        header.setFont(QFont('High Tower Text', 18, QFont.Weight.Bold))
        header.setStyleSheet("color: #176276;")
        header.setAlignment(Qt.AlignmentFlag.AlignCenter)

        # Table Widget to Display Structure
        self.structure_table = QTableWidget()
        self.structure_table.setFont(QFont('Gadugi', 12))
        self.structure_table.setColumnCount(len(columns))
        self.structure_table.setHorizontalHeaderLabels(columns)
        self.structure_table.setRowCount(len(structure_data))
        for row_idx, row_data in enumerate(structure_data):
            for col_idx, col_data in enumerate(row_data):
                self.structure_table.setItem(row_idx, col_idx, QTableWidgetItem(str(col_data)))
        self.structure_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)

        # Export Button
        export_btn = QPushButton("EXPORT TABLE STRUCTRE")
        export_btn.setFont(QFont('Candara', 14))
        export_btn.setStyleSheet("background-color: #176276; color: #F9F4E0;")
        export_btn.clicked.connect(self.export_structure)

        # Add Widgets to Layout
        layout.addWidget(header)
        layout.addWidget(self.structure_table)
        layout.addWidget(export_btn)
        layout.setAlignment(Qt.AlignmentFlag.AlignTop)

    def export_structure(self):
        """
        Export the structure as a file (e.g., PDF, Word).
        """
        options = ["PDF", "Word", "Excel"]
        file_type, ok = QInputDialog.getItem(self, "Export Structure", "Select export format:", options, 0, False)
        if ok and file_type:
            file_path, _ = QFileDialog.getSaveFileName(self, "Save File", f"{self.table_name}_structure.{file_type.lower()}",
                                                       f"{file_type} Files (*.{file_type.lower()})")
            if file_path:
                try:
                    if file_type == "PDF":
                        printer = QPrinter(QPrinter.PrinterMode.HighResolution)
                        printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
                        printer.setOutputFileName(file_path)
                        print_dialog = QPrintDialog(printer)
                        if print_dialog.exec() == QPrintDialog.DialogCode.Accepted:
                            print_dialog.paintRequested.emit(printer)
                    elif file_type == "Word":
                        from docx import Document
                        doc = Document()
                        doc.add_heading(f"Structure of Table: {self.table_name}", level=1)
                        for row in self.structure_data:
                            doc.add_paragraph(str(row))
                        doc.save(file_path)
                    elif file_type == "Excel":
                        import pandas as pd
                        pd.DataFrame(self.structure_data, columns=self.columns).to_excel(file_path, index=False)
                    QMessageBox.information(self, "Success", f"Structure exported successfully to {file_path}!")
                except Exception as e:
                    QMessageBox.critical(self, "Error", f"Failed to export structure: {str(e)}")

class TableViewer(QWidget):
    def __init__(self, connection, database_name, table_name, main_window):
        super().__init__()
        self.connection = connection
        self.database_name = database_name
        self.table_name = table_name
        self.main_window = main_window
        self.original_data = []  # Stores the original data for refresh
        self.filtered = False  # Tracks if data is filtered or sorted

        self.setWindowTitle(f'{table_name} - {database_name}')
        self.setStyleSheet("background-color: white; color: black;")
        layout = QVBoxLayout()
        self.setLayout(layout)

        # Header
        header = QLabel(f"TABLE: {table_name}  |  DATABASE: {database_name}")
        header.setFont(QFont('High Tower Text', 24, QFont.Weight.Bold))
        header.setStyleSheet("color: #176276;")
        header.setAlignment(Qt.AlignmentFlag.AlignCenter)

        # Toolbar and Table Widget
        self.tool_bar = TableContentToolBar(self.connection, self)
        self.table_widget = QTableWidget()
        self.table_widget.setFont(QFont('Gadugi', 12))

        # Add Widgets to Layout
        layout.addWidget(header)
        layout.addWidget(self.tool_bar)
        layout.addWidget(self.table_widget)
        layout.setAlignment(Qt.AlignmentFlag.AlignTop)

        # Populate table data
        self.update_table()

    def update_table(self):
        """
        Fetch and display table content from the database.
        """
        try:
            cursor = self.connection.cursor(dictionary=True)
            cursor.execute(f"USE {self.database_name}")
            cursor.execute(f"SELECT * FROM {self.table_name}")
            self.original_data = cursor.fetchall()
            cursor.close()

            # Update the table display
            self.refresh_table(self.original_data)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to fetch table data: {str(e)}")

    def refresh_table(self, data):
        """
        Updates the QTableWidget with new data.
        """
        self.table_widget.clear()
        if data:
            headers = data[0].keys()
            self.table_widget.setColumnCount(len(headers))
            self.table_widget.setHorizontalHeaderLabels(headers)
            self.table_widget.setRowCount(len(data))
            for row_idx, row_data in enumerate(data):
                for col_idx, col_data in enumerate(row_data.values()):
                    self.table_widget.setItem(row_idx, col_idx, QTableWidgetItem(str(col_data)))
            self.table_widget.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        else:
            self.table_widget.setRowCount(0)

    def restore_original_content(self):
        """
        Restores the original table content when the reset or refresh button is clicked.
        """
        if self.filtered:
            self.refresh_table(self.original_data)
            self.filtered = False


class TableContentToolBar(QWidget):
    def __init__(self, connection, parent):
        super().__init__()
        self.connection = connection
        self.parent = parent
        self.filtered = False  # Tracks if the table is filtered or sorted
        self.setStyleSheet("background-color: white; color: black;")
        layout = QHBoxLayout()
        self.setLayout(layout)

        # Sort Button
        sort_btn = QPushButton("SORT TABLE CONTENT")
        sort_btn.setFont(QFont('Candara', 14))
        sort_btn.setStyleSheet("background-color: #176276; color: #F9F4E0;")
        sort_btn.clicked.connect(self.sort_table)

        # Search Button
        search_btn = QPushButton("SEARCH TABLE CONTENT")
        search_btn.setFont(QFont('Candara', 14))
        search_btn.setStyleSheet("background-color: #176276; color: #F9F4E0;")
        search_btn.clicked.connect(self.search_table_content)

        # Export Button
        export_btn = QPushButton("EXPORT TABLE")
        export_btn.setFont(QFont('Candara', 14))
        export_btn.setStyleSheet("background-color: #176276; color: #F9F4E0;")
        export_btn.clicked.connect(self.export_table)

        # Visualize Button
        visualize_btn = QPushButton("VISUALIZE TABLE")
        visualize_btn.setFont(QFont('Candara', 14))
        visualize_btn.setStyleSheet("background-color: #176276; color: #F9F4E0;")
        visualize_btn.clicked.connect(self.visualize_table)

        # Reset Button
        self.reset_btn = QPushButton("RESET")
        self.reset_btn.setFont(QFont('Candara', 14))
        self.reset_btn.setStyleSheet("background-color: #176276; color: #F9F4E0;")
        self.reset_btn.clicked.connect(self.reset_table)
        self.reset_btn.setVisible(False)

        # Go Back Button
        go_back_btn = QPushButton("GO BACK")
        go_back_btn.setFont(QFont('Candara', 14))
        go_back_btn.setStyleSheet("background-color: #176276; color: #F9F4E0;")
        go_back_btn.clicked.connect(self.go_back)

        # Add Buttons to Layout
        layout.addWidget(sort_btn)
        layout.addWidget(search_btn)
        layout.addWidget(export_btn)
        layout.addWidget(visualize_btn)
        layout.addWidget(self.reset_btn)
        layout.addWidget(go_back_btn)
        layout.setAlignment(Qt.AlignmentFlag.AlignTop)

    def sort_table(self):
        """
        Sort the table based on user-selected criteria.
        """
        columns = list(self.parent.original_data[0].keys())
        sort_column, ok = QInputDialog.getItem(self, "Sort Table", "Select column to sort by:", columns, 0, False)
        if ok and sort_column:
            criteria, ok = QInputDialog.getItem(
                self, "Sort Criteria", f"Select sorting order for column '{sort_column}':",
                ["Ascending", "Descending"], 0, False
            )
            if ok:
                reverse = criteria == "Descending"
                sorted_data = sorted(self.parent.original_data, key=lambda x: x[sort_column], reverse=reverse)
                self.parent.refresh_table(sorted_data)
                self.filtered = True
                self.reset_btn.setVisible(True)

    def search_table_content(self):
        """
        Search for a keyword in the table content.
        """
        keyword, ok = QInputDialog.getText(self, "Search Table Content", "Enter keyword to search:")
        if ok and keyword:
            filtered_data = [row for row in self.parent.original_data if keyword.lower() in str(row).lower()]
            self.parent.refresh_table(filtered_data)
            self.filtered = True
            self.reset_btn.setVisible(True)

    def export_table(self):
        """
        Export table content to various formats.
        """
        options = ["Excel", "CSV", "PDF", "Word"]
        file_type, ok = QInputDialog.getItem(self, "Export Table", "Select file format:", options, 0, False)
        if ok and file_type:
            file_path, _ = QFileDialog.getSaveFileName(self, "Save File", f"{self.parent.table_name}.{file_type.lower()}")
            if file_path:
                try:
                    import pandas as pd
                    df = pd.DataFrame(self.parent.original_data)
                    if file_type == "Excel":
                        df.to_excel(file_path, index=False)
                    elif file_type == "CSV":
                        df.to_csv(file_path, index=False)
                    elif file_type == "PDF":
                        from fpdf import FPDF
                        pdf = FPDF()
                        pdf.add_page()
                        pdf.set_font("Arial", size=12)
                        for i, row in enumerate(df.values.tolist()):
                            pdf.cell(200, 10, txt=str(row), ln=i + 1)
                        pdf.output(file_path)
                    elif file_type == "Word":
                        from docx import Document
                        doc = Document()
                        doc.add_heading(f"Table: {self.parent.table_name}", level=1)
                        for row in df.values.tolist():
                            doc.add_paragraph(str(row))
                        doc.save(file_path)
                    QMessageBox.information(self, "Success", f"Table exported successfully to {file_path}!")
                except Exception as e:
                    QMessageBox.critical(self, "Error", f"Failed to export table: {str(e)}")


    def visualize_table(self):
        """
        Visualizes the table data intelligently with proper axis labels and column names.
        """
        try:
            import matplotlib.pyplot as plt
            import seaborn as sns
            import pandas as pd

            # Columns available in the table
            columns = list(self.parent.original_data[0].keys())

            # Prompt user to select columns for visualization
            selected_columns, ok = QInputDialog.getMultiLineText(
                self, "Select Columns", "Enter columns to visualize (comma-separated):", ", ".join(columns)
            )
            if not ok or not selected_columns:
                return

            selected_columns = [col.strip() for col in selected_columns.split(",")]
            df = pd.DataFrame(self.parent.original_data)[selected_columns]

            # Analyze data types of the selected columns
            data_types = {col: type(next((row[col] for row in self.parent.original_data if row[col] is not None), None))
                          for col in selected_columns}
            numeric_columns = [col for col, dtype in data_types.items() if dtype in [int, float]]

            # Validate columns for visualization
            if len(selected_columns) < 1:
                QMessageBox.warning(self, "Visualization Error", "Please select at least one column.")
                return

            # Prompt user for chart type
            chart_type, ok = QInputDialog.getItem(
                self, "Visualize Table", "Select chart type:",
                ["Bar", "Pie", "Line", "Scatter", "Heatmap"], 0, False
            )
            if not ok:
                return

            plt.figure(figsize=(10, 6))  # Initialize the figure

            # Generate visualization based on the selected chart type
            if chart_type == "Bar":
                ax = df.plot(kind="bar", x=selected_columns[0], y=selected_columns[1] if len(selected_columns) > 1 else None)
                ax.set_title(f"{chart_type} Chart")
                ax.set_xlabel(selected_columns[0])  # x-axis label
                ax.set_ylabel(selected_columns[1] if len(selected_columns) > 1 else "Values")  # y-axis label
            elif chart_type == "Pie":
                if len(numeric_columns) > 0:
                    ax = df.plot(
                        kind="pie", y=numeric_columns[0], autopct='%1.1f%%', labels=df[selected_columns[0]],
                        title=f"{chart_type} Chart"
                    )
                    ax.set_ylabel("")  # Remove the default ylabel
                else:
                    QMessageBox.warning(self, "Visualization Error", "Pie chart requires numeric columns for values.")
                    return
            elif chart_type == "Line":
                ax = df.plot(kind="line", x=selected_columns[0], y=selected_columns[1] if len(selected_columns) > 1 else None)
                ax.set_title(f"{chart_type} Graph")
                ax.set_xlabel(selected_columns[0])  # x-axis label
                ax.set_ylabel(selected_columns[1] if len(selected_columns) > 1 else "Values")  # y-axis label
            elif chart_type == "Scatter":
                if len(numeric_columns) >= 2:
                    ax = sns.scatterplot(data=df, x=numeric_columns[0], y=numeric_columns[1])
                    ax.set_title(f"{chart_type} Plot ({numeric_columns[0]} vs {numeric_columns[1]})")
                    ax.set_xlabel(numeric_columns[0])  # x-axis label
                    ax.set_ylabel(numeric_columns[1])  # y-axis label
                else:
                    QMessageBox.warning(self, "Visualization Error", "Scatter plot requires at least two numeric columns.")
                    return
            elif chart_type == "Heatmap":
                if len(numeric_columns) >= 2:
                    sns.heatmap(df[numeric_columns].corr(), annot=True, cmap="coolwarm", xticklabels=numeric_columns,
                                yticklabels=numeric_columns)
                    plt.title(f"{chart_type} of Correlation Matrix")
                else:
                    QMessageBox.warning(self, "Visualization Error", "Heatmap requires at least two numeric columns.")
                    return

            # Set window title dynamically
            plt.gcf().canvas.manager.set_window_title(
                f"MENSA - VISUALIZE TABLE {self.parent.table_name} as {chart_type}"
            )

            # Show the plot
            plt.tight_layout()
            plt.show()

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Visualization failed: {str(e)}")


    def reset_table(self):
        """
        Resets the table to its original state.
        """
        self.parent.restore_original_content()
        self.reset_btn.setVisible(False)
        self.filtered = False

    def go_back(self):
        """
        Navigates back to the Database Arena.
        """
        self.parent.main_window.switch_widget(self.parent.main_window.db_arena)

class TableForm(QWidget):
    def __init__(self, connection, database_name, action, table_name=None):
        super().__init__()
        self.connection = connection
        self.database_name = database_name
        self.action = action
        self.table_name = table_name
        self.setWindowTitle(f"{action.capitalize()} Table - {database_name}")
        self.setStyleSheet("background-color: white; color: black;")
        self.setFixedSize(400, 300)
        layout = QVBoxLayout()
        self.setLayout(layout)

        # Form Layout
        form_layout = QFormLayout()
        self.table_name_input = QLineEdit()
        if self.action == "rename" and self.table_name:
            self.table_name_input.setText(self.table_name)
        form_layout.addRow("Table Name:", self.table_name_input)

        self.columns_input = QLineEdit()
        form_layout.addRow("Columns (name type, name type):", self.columns_input)

        # Submit Button
        submit_btn = QPushButton("Submit")
        submit_btn.setFont(QFont('High Tower Text', 14))
        submit_btn.setStyleSheet("background-color: #2d388a; color: #00aeef;")
        submit_btn.clicked.connect(self.handle_submit)

        # Add Widgets to Layout
        layout.addLayout(form_layout)
        layout.addWidget(submit_btn)

    def handle_submit(self):
        """
        Handle table creation or renaming.
        """
        try:
            table_name = self.table_name_input.text()
            columns = self.columns_input.text()
            cursor = self.connection.cursor()
            cursor.execute(f"USE {self.database_name}")
            if self.action == "create":
                cursor.execute(f"CREATE TABLE {table_name} ({columns})")
            elif self.action == "rename" and self.table_name:
                cursor.execute(f"ALTER TABLE {self.table_name} RENAME TO {table_name}")
            self.connection.commit()
            cursor.close()
            QMessageBox.information(self, "Success", f"Table {self.action}d successfully!")
            self.close()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to {self.action} table: {str(e)}")

if __name__ == '__main__':
    app = QApplication(sys.argv)
    loading_screen = LoadingScreen()
    loading_screen.show()
    sys.exit(app.exec())
